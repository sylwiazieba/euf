from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('🌱 EUF Harvest Log', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Easton Urban Farm — Harvest Tracking Form')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(11)
subtitle_format.font.italic = True

# Add instructions
doc.add_heading('Instructions:', level=2)
instructions = doc.add_paragraph()
instructions.add_run('Units: ').bold = True
instructions.add_run('lbs (pounds), bunch (herbs), bin (full harvested), pot (individual seedlings)\n')
instructions.add_run('Destination: ').bold = True
instructions.add_run('Leave blank if going to Food Pantry. Otherwise, write organization name (New Bethany, Pembroke, Project of Easton, Safe Harbor, etc.)\n')
instructions.add_run('Variety: ').bold = True
instructions.add_run('e.g., "Cherry", "Roma", "Basil", or leave blank if unknown\n')
instructions.add_run('Notes: ').bold = True
instructions.add_run('Any special comments, issues, or observations\n')

doc.add_paragraph()

# Create table
table = doc.add_table(rows=21, cols=6)
table.style = 'Light Grid Accent 1'

# Set header row
header_cells = table.rows[0].cells
headers = ['Date\n(MM/DD)', 'Crop Name', 'Variety', 'Amount & Unit\n(e.g., "50 lbs")', 'Destination\n(blank = Food Pantry)', 'Notes/Comments']

for i, header_text in enumerate(headers):
    cell = header_cells[i]
    cell.text = header_text
    # Format header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell._element.get_or_add_tcPr().append(
        doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd',
                                 {'{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': '2E8B57'})
    )

# Set column widths
table.autofit = False
table.allow_autofit = False
widths = [Inches(1.0), Inches(1.3), Inches(1.1), Inches(1.3), Inches(1.5), Inches(1.8)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

# Save document
doc.save('harvest_log_template.docx')
print("✅ Created: harvest_log_template.docx")
